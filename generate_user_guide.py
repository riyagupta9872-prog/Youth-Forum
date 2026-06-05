"""
Generates USER_GUIDE.docx for the Sakhi Sang Attendance Tracker.
Run once: python generate_user_guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
GREEN       = RGBColor(0x1a, 0x5c, 0x3a)   # primary brand green
DARK        = RGBColor(0x1a, 0x1a, 0x2e)
ACCENT_BLUE = RGBColor(0x0d, 0x6e, 0xfd)
TIP_BG      = RGBColor(0xe8, 0xf5, 0xe9)
WARN_BG     = RGBColor(0xff, 0xf3, 0xcd)
LIGHT_GRAY  = RGBColor(0xf5, 0xf5, 0xf5)
MID_GRAY    = RGBColor(0xcc, 0xcc, 0xcc)
WHITE       = RGBColor(0xff, 0xff, 0xff)
HEAD_BG     = RGBColor(0x1a, 0x5c, 0x3a)

# ── Helper: shading on a table cell ──────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ──────────────────────────────────────────────────
def set_table_borders(table, color="CCCCCC", size="4"):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

# ── Helper: paragraph formatting shortcuts ────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = GREEN
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:color'), '1a5c3a')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = GREEN
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return p

def body(text, bold_parts=None):
    """Normal paragraph. bold_parts = list of substrings to bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                if idx > 0:
                    r = p.add_run(remaining[:idx])
                    r.font.size = Pt(11)
                rb = p.add_run(bp)
                rb.bold = True
                rb.font.size = Pt(11)
                remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def callout(label, text, bg_hex, border_hex):
    """A shaded callout box (Tip / Warning)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg_hex)
    # border colour
    set_table_borders(tbl, border_hex, "6")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{label}  ")
    r1.bold           = True
    r1.font.size      = Pt(11)
    r1.font.color.rgb = RGBColor(int(border_hex[:2],16),
                                  int(border_hex[2:4],16),
                                  int(border_hex[4:],16))
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()  # spacing after callout

def tip(text):
    callout("💡 Tip:", text, "e8f5e9", "1a5c3a")

def warn(text):
    callout("⚠️ Warning:", text, "fff3cd", "856404")

def screenshot(desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ Screenshot: {desc} ]")
    r.italic          = True
    r.font.size       = Pt(10)
    r.font.color.rgb  = RGBColor(0x88, 0x88, 0x88)
    # dashed box via shading table trick
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade_cell(cell, "f9f9f9")
    set_table_borders(tbl, "aaaaaa", "4")
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"📸  {desc}")
    r2.italic         = True
    r2.font.size      = Pt(10)
    r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()

def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl, "cccccc", "4")
    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], "1a5c3a")
        hdr_cells[i].paragraphs[0].clear()
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri+1].cells
        bg = "ffffff" if ri % 2 == 0 else "f5f5f5"
        for ci, val in enumerate(row):
            shade_cell(cells[ci], bg)
            cells[ci].paragraphs[0].clear()
            if isinstance(val, tuple):
                txt, bold = val
            else:
                txt, bold = val, False
            run = cells[ci].paragraphs[0].add_run(str(txt))
            run.font.size = Pt(10)
            run.bold = bold
        if col_widths:
            for ci, w in enumerate(col_widths):
                cells[ci].width = Inches(w)
    # apply col widths to header too
    if col_widths:
        for ci, w in enumerate(col_widths):
            hdr_cells[ci].width = Inches(w)
    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
shade_cell(cover.cell(0,0), "1a5c3a")
set_table_borders(cover, "1a5c3a", "0")
cell = cover.cell(0,0)
cell.width = Inches(6)

p = cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("Sakhi Sang")
r.bold           = True
r.font.size      = Pt(32)
r.font.color.rgb = WHITE

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Attendance Tracker")
r2.bold           = True
r2.font.size      = Pt(24)
r2.font.color.rgb = RGBColor(0xc8, 0xe6, 0xc9)

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Complete User Guide")
r3.font.size      = Pt(14)
r3.font.color.rgb = RGBColor(0xc8, 0xe6, 0xc9)
r3.italic = True

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(40)
r4 = p4.add_run("Plain English • Step-by-Step • All Roles")
r4.font.size      = Pt(11)
r4.font.color.rgb = RGBColor(0xa5, 0xd6, 0xa7)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    ("1.", "System Overview"),
    ("2.", "Getting Started — Login & Signup"),
    ("3.", "Who Can Do What — Access Table"),
    ("4.", "Where Everything Lives — Interface Map"),
    ("5.", "Feature Deep-Dive"),
    ("  5A.", "Home Dashboard"),
    ("  5B.", "Managing Devotees"),
    ("  5C.", "Calling (Weekly Follow-Up)"),
    ("  5D.", "Attendance Marking"),
    ("  5E.", "Activities — Books, Service, Registration, Donation"),
    ("  5F.", "Reports & Analytics"),
    ("  5G.", "Care Tab — Following Up on Absent Devotees"),
    ("  5H.", "Events"),
    ("  5I.", "Calling Management (Super Admin)"),
    ("6.", "The Filter Bar — Controlling What You See"),
    ("7.", "Admin Panel — Managing Users & Sessions"),
    ("8.", "Excel Import & Export"),
    ("9.", "Tips, Warnings & Common Questions"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(11)
    if not num.startswith(" "):
        r.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. System Overview")
body("The Sakhi Sang Attendance Tracker is a web app built for the coordinators and facilitators of the Sakhi Sang women's devotee group. It brings together everything you need to manage your group in one place:")
for line in [
    "Keep a database of all devotees — their contact details, team, spiritual practices, family background, and more.",
    "Track weekly calling — who called whom, who said they would come, and whether they actually came.",
    "Mark attendance at Sunday sessions live on your phone or computer.",
    "Record activities throughout the week — books distributed, donations collected, registrations, and service.",
    "See reports and charts to understand how your team is performing.",
    "Identify who needs care — devotees who are absent, inactive, or said they would come but didn't.",
]:
    bullet(line)

body("")
body("Everything is organised by team (for example: Champaklata, Lalita, Rangadevi). Each coordinator manages their own team, while Super Admins can see everything across all teams.")
tip("This app works on your phone browser too — no installation needed. Just open the link and log in.")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Getting Started — Login & Signup")

h2("Signing Up for the First Time")
for step in [
    "Open the app link in your browser.",
    "Click \"Sign Up\" on the login screen.",
    "Enter your name, email address, and a password.",
    "Click \"Request Access\".",
    "You will see a message: \"Your request is awaiting approval\" — this means a Super Admin needs to approve your account before you can log in.",
]:
    numbered(step)
warn("You cannot use the app until a Super Admin approves you. If you've been waiting too long, contact your Super Admin directly.")

h2("Logging In")
for step in [
    "Open the app link.",
    "Enter your email and password.",
    "Click \"Login\".",
    "The app will open to your Home / Dashboard tab.",
]:
    numbered(step)

h2("If Your Account Was Rejected")
body("You will see the message: \"This account was not approved.\" Contact your Super Admin to find out why and resubmit if needed.")
screenshot("Login screen with email, password fields and Login / Sign Up buttons")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — ACCESS CONTROL MATRIX
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Who Can Do What — Access Table")
body("There are three types of accounts in this app:")

simple_table(
    ["What You Want to Do", "Super Admin", "Coordinator", "Facilitator / Att. Volunteer"],
    [
        ("See all teams' data",                        "✅", "❌ own team only", "❌ own team only"),
        ("Add / edit devotees",                        "✅", "✅", "❌"),
        ("Mark devotee as Not Interested or remove",   "✅", "❌", "❌"),
        ("Weekly calling & submit calling status",     "✅", "✅", "❌"),
        ("Mark attendance at Sunday sessions",         "✅", "✅", "✅ (if enabled)"),
        ("Log books, donations, registrations, service","✅", "✅", "❌"),
        ("View reports and care alerts",               "✅", "✅", "❌"),
        ("Manage events",                              "✅", "✅", "❌"),
        ("Approve / reject new user signups",          "✅", "❌", "❌"),
        ("Configure Sunday sessions",                  "✅", "❌", "❌"),
        ("Clear / delete data",                        "✅", "❌", "❌"),
        ("Manage all users' roles",                    "✅", "❌", "❌"),
        ("View Calling Management grid",               "✅", "❌", "❌"),
    ],
    col_widths=[3.2, 1.0, 1.2, 1.7]
)
tip("Super Admins can grant an Attendance Volunteer access to mark attendance across all teams by enabling \"Att. Seva\" in the Admin Panel.")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — INTERFACE MAP
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Where Everything Lives — Interface Map")
body("When you log in, you see a row of tabs at the top of the screen (or at the bottom on a phone). Here is what each tab does:")

simple_table(
    ["Tab Name", "What It Does"],
    [
        ("Home / Dashboard",   "Quick summary stats + fast-entry shortcuts for today's activities"),
        ("Devotees",           "Full database of all devotees — add, edit, search, view profiles"),
        ("Calling",            "Weekly calling tracker — record who you called, whether they're coming, and why not"),
        ("Attendance",         "Mark who came on Sunday; view past attendance sheets"),
        ("Books",              "Log books distributed this week"),
        ("Service",            "Log service done by devotees"),
        ("Registration",       "Log event or programme registrations"),
        ("Donation",           "Log donations received"),
        ("Care",               "Alerts for absent, inactive, or at-risk devotees"),
        ("Events",             "Create and manage special events; track who attended"),
        ("Calling Mgmt",       "Super Admin only — overview of all teams' calling progress"),
    ],
    col_widths=[1.5, 5.5]
)
screenshot("Top navigation bar showing all tab names")

h2("The Filter Bar (Below the Tabs)")
body("Just below the tab row, there is a Filter Bar with three buttons:")
bullet("Session — which Sunday you are looking at (defaults to the latest configured Sunday)")
bullet("Team — which team's data to show (Coordinators are locked to their own team)")
bullet("Calling By — filter by which coordinator is doing the calling")
body("")
body("Changing any of these filters will instantly update the data in whichever tab you are on.")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — FEATURE DEEP-DIVE
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Feature Deep-Dive")

# ── 5A HOME ──────────────────────────────────────────────────────────────────
h2("5A. Home Dashboard")
body("The Home tab is your quick-action hub. It shows you a greeting (\"Hare Krishna, [Your Name]!\") and tiles you can tap to do common tasks without navigating away.")

h3("Quick-Entry Buttons")
simple_table(
    ["Button", "What It Opens"],
    [
        ("Attendance Report",  "A summary table of today's calling and attendance numbers per team"),
        ("Book Distribution",  "A form to quickly log books given out today"),
        ("Donation",           "A form to log a donation received"),
        ("Registration",       "A form to log a registration"),
        ("Service",            "A form to log service done by a devotee"),
    ],
    col_widths=[1.8, 5.2]
)

h3("How to Log a Book Entry from Home")
for step in [
    "Tap the Book Distribution button on the Home tab. A panel slides up from the bottom.",
    "In the Devotee field, type the devotee's name or mobile number to search. Select them from the list.",
    "The Team field will fill in automatically.",
    "Set the Date (defaults to today).",
    "Enter the Quantity of books.",
    "Tap \"Save Entry\".",
]:
    numbered(step)
tip("The same steps apply for Service, Registration, and Donation — just tap the relevant button.")
tip("The Attendance Report quick button shows you a real-time snapshot of calling vs attendance per team — useful before the Sunday session starts.")

# ── 5B DEVOTEES ───────────────────────────────────────────────────────────────
h2("5B. Managing Devotees")
body("The Devotees tab is where all devotee information is stored and managed.")

h3("Searching for a Devotee")
for step in [
    "Go to the Devotees tab.",
    "Use the search box at the top to type a name or mobile number.",
    "The list filters instantly as you type.",
    "You can also filter by status (Most Serious, Serious, New Devotee, Inactive, etc.) using the status filter.",
]:
    numbered(step)
screenshot("Devotee list with search box and status filter dropdown")

h3("Adding a New Devotee")
numbered("Go to the Devotees tab.")
numbered("Click the \"+ Add Devotee\" button.")
numbered("A form opens with 5 sections:")

sections = [
    ("Section 1 — Personal Identity",
     ["Full Name (required)", "Mobile Number (required, 10 digits)", "Alternate Mobile",
      "Residential Address", "Date of Birth", "Email"]),
    ("Section 2 — Team & Status",
     ["Team (select from dropdown)", "Devotee Status: Most Serious / Serious / Expected to be Serious / New Devotee / Inactive",
      "Date of Joining", "Reference By — who introduced them (search from existing devotees)",
      "Facilitator — who takes care of them",
      "Calling By — who calls them each week (must have a system login)"]),
    ("Section 3 — Professional Profile",
     ["Education / Qualification", "Profession / Occupation"]),
    ("Section 4 — Sadhana & Practices",
     ["Daily Chanting Rounds", "Reading and Hearing habits",
      "Wears Tilak (Yes/No), Kanthi (Yes/No), Gopi Dress (Yes/No)",
      "Plays Instrument — if Yes, enter instrument name",
      "Attends Kirtan Classes (Yes/No)"]),
    ("Section 5 — Social & Family",
     ["Total Family Members", "Family Members attending class",
      "Family's Attitude Towards Devotion", "Hobbies & Interests"]),
]
for sec_title, fields in sections:
    h3(f"   {sec_title}")
    for f in fields:
        bullet(f, level=1)

numbered("Fill in at least the required fields (Name, Mobile) and click \"Save\".")
tip("A colour circle on the profile shows how complete the profile is — red means less than 50% filled, green means 80%+. Try to fill as much as possible.")

h3("Editing a Devotee's Profile")
for step in [
    "Click on the devotee's card in the list to open their profile.",
    "Click the \"Edit\" button.",
    "Update any field and click \"Save\".",
]:
    numbered(step)
warn("Only Coordinators and Super Admins can edit profiles. All changes are tracked in an audit history — nothing is ever silently overwritten.")

h3("Marking a Devotee as \"Not Interested\"")
body("Use this for devotees who no longer wish to be contacted.")
for step in [
    "Open the devotee's profile.",
    "Click \"Not Interested\" (only Super Admins can see this button).",
    "Confirm the action.",
    "The devotee moves to a separate list and is excluded from all calling counts.",
]:
    numbered(step)
warn("This does NOT delete the devotee. Their record is preserved. A Super Admin can reverse this if needed.")

# ── 5C CALLING ────────────────────────────────────────────────────────────────
h2("5C. Calling (Weekly Follow-Up)")
body("The Calling tab is where you record the outcome of your weekly calls to devotees.")

h3("Understanding the Calling Tab")
body("Each week, before Sunday's session, coordinators call each devotee in their list and record two things:")
bullet("Are they coming this Sunday? (Yes / No)")
bullet("If no, why?")
body("The calling window opens on a specific day (set by the Super Admin) and closes at 9 PM on the calling deadline.")
screenshot("Calling tab showing devotee list with Coming toggle and Reason dropdown")

h3("How to Update Calling Status")
for step in [
    "Go to the Calling tab.",
    "You'll see your list of devotees.",
]:
    numbered(step)
numbered("For each devotee:")
bullet('If they said "Yes, I\'m coming" -> tap "Mark Yes". It turns green and shows "Coming".', level=1)
bullet("If they are not coming, click the row and select a Reason:", level=1)
reasons = ["Did not pick call", "Incoming not available", "Wrong number",
           "Out of station (enter the date they return)", "Exams (enter the date they return)",
           "Shifted to online class", "Festival Calling", "Not Interested (this week)"]
for r in reasons:
    bullet(r, level=2)
numbered("You can also add a short note in the Notes field.")
tip("You can see at the top of the tab how many are Confirmed, Not Reached, Unavailable, etc. These update in real time.")

h3("Submitting Your Calling")
body("Once you've called everyone, you must officially submit your calling:")
for step in [
    "At the bottom of the Calling tab, look for: \"Done calling for [DATE]? [Submit Calling]\"",
    "Click \"Submit Calling\".",
    "You'll see a confirmation message with a timestamp.",
]:
    numbered(step)
warn("Submit before 9 PM on the deadline. If you submit after 9 PM, it will be marked as \"Late\" in reports. The initial submission time is locked — it won't change if you update calling status later.")

h3("Viewing Calling Reports")
body("There are two report sub-tabs inside Calling:")
bullet("Weekly Report — Shows each team's calling summary (total called, said yes, actually came). Click \"Export\" to download as Excel.")
bullet("Accuracy Report — Shows how many devotees said they would come but didn't attend. Useful for improving calling quality.")

# ── 5D ATTENDANCE ─────────────────────────────────────────────────────────────
h2("5D. Attendance Marking")
body("The Attendance tab is used on Sunday to mark who is present at the session.")

h3("Before the Session — What the Super Admin Does")
body("A Super Admin must configure the session first:")
for step in [
    "Go to the Attendance tab.",
    "Click \"Configure Sunday\" (or the settings gear icon).",
    "Enter the Calling Date, the Session / Attendance Date, Topic (optional), Speaker Name (optional), and Session Type (Regular or Festival).",
    "Save. The session is now active.",
]:
    numbered(step)

h3("Marking Attendance Live (on Sunday)")
for step in [
    "Go to the Attendance tab.",
    "You'll see a list of devotees — those who said \"Yes\" on calling appear with a \"Confirmed\" badge.",
    "As each devotee arrives, find their card and tap \"Present\".",
    "The button changes to show \"P [Time]\" — the time they were marked.",
    "If you marked someone by mistake, tap \"Undo\".",
]:
    numbered(step)

body("At the top, you'll see 4 live counters updating:")
simple_table(
    ["Counter", "What It Means"],
    [
        ("Total Confirmed", "How many said they would come"),
        ("Present",         "How many have been marked so far"),
        ("New",             "First-timers today"),
        ("Total Present",   "Combined present count"),
    ],
    col_widths=[1.8, 5.2]
)
screenshot("Attendance tab with devotee cards showing Present button and live counters")
tip("If a new devotee arrives who is not in the system, you can register them on the spot — there is an option to add a new devotee directly from the Attendance tab.")

h3("Viewing Past Attendance")
for step in [
    "Change the Session in the Filter Bar to a past date.",
    "The Attendance tab will switch to a historical view showing a full attendance sheet.",
    "The sheet is colour-coded: light blue = 30+ sessions attended, light green = 15+ sessions, light yellow = 5+ sessions.",
]:
    numbered(step)

# ── 5E ACTIVITIES ─────────────────────────────────────────────────────────────
h2("5E. Activities — Books, Service, Registration, Donation")
body("Each of these four tabs works the same way — they have a Log Entry section and a Reports section.")

h3("Logging a New Entry")
numbered("Go to the relevant tab (e.g., Books).")
numbered("You'll see a Log Entry section at the top.")
numbered("Fill in:")
bullet("Devotee — search and select the devotee (not required for Donation)", level=1)
bullet("Team — auto-fills from the devotee; can be changed", level=1)
bullet("Date — defaults to today", level=1)
bullet("Books → Quantity; Service → Description; Registration → Count; Donation → Amount in ₹ + optional Note", level=1)
numbered("Click \"Save Entry\".")
numbered("The entry appears in the Recent Entries list below.")

h3("Viewing Reports")
for step in [
    "Scroll down (or click the \"Reports\" sub-tab).",
    "Set the From and To dates for the date range you want.",
    "Click \"Refresh\" to update.",
    "You'll see summary tiles (Grand Total, Number of Entries, Date Range) and a table broken down by team — click a team row to expand and see individual entries.",
    "Click \"Export Excel\" to download the report.",
]:
    numbered(step)

# ── 5F REPORTS ────────────────────────────────────────────────────────────────
h2("5F. Reports & Analytics")

h3("Dashboard Tiles (Home Tab)")
body("At the top of the Home tab, six tiles show this week's numbers at a glance:")
simple_table(
    ["Tile", "What It Shows"],
    [
        ("Attended",          "How many came out of the calling list"),
        ("Calling Accuracy",  "Percentage of \"Yes\" responses that actually attended"),
        ("Books",             "Books distributed this week"),
        ("Services",          "Services logged this week"),
        ("Registrations",     "Registrations logged this week"),
        ("Donation",          "Total donations collected (₹)"),
    ],
    col_widths=[1.8, 5.2]
)

h3("Coordinator Performance Grid")
body("The Reports tab shows a detailed grid of each team / coordinator. Click on any number to see the list of devotees behind that number.")
simple_table(
    ["Column", "Meaning"],
    [
        ("Team",       "Team name"),
        ("Called",     "How many devotees were called"),
        ("Yes",        "How many said they would come"),
        ("Came",       "How many actually attended"),
        ("Target",     "Attendance target set for the team"),
        ("%",          "Achievement percentage"),
        ("Books",      "Books logged"),
        ("Service",    "Service entries"),
        ("Reg.",       "Registrations"),
        ("Donation ₹", "Total donation amount"),
    ],
    col_widths=[1.5, 5.5]
)

h3("Attendance Reports — Period Analysis")
for step in [
    "Go to the Attendance tab and look for the Reports section.",
    "Choose a period: Single Session, Month, Quarter, or FY (April to March).",
    "Select a sub-tab for the view you need:",
]:
    numbered(step)
sub_tabs = [
    ("Attendance Sheet", "Full per-devotee grid showing calling status and attendance for each session"),
    ("Late Comers",      "Devotees who arrived after a certain time"),
    ("New Comers",       "First-time attendees"),
    ("Serious Analysis", "Breakdown by devotee status (Most Serious, Serious, etc.)"),
    ("Team Leaderboard", "Ranks teams by attendance achievement (🥇 🥈 🥉 medals)"),
    ("Trends",           "A line chart of attendance numbers over time"),
]
for name, desc in sub_tabs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
screenshot("Team Leaderboard tab with medal rankings")

# ── 5G CARE ───────────────────────────────────────────────────────────────────
h2("5G. Care Tab — Following Up on Absent Devotees")
body("The Care tab is a set of automatic alerts to help you follow up with devotees who need attention. It shows 5 alert cards based on the session selected in the Filter Bar:")

simple_table(
    ["Alert Card", "Who It Shows"],
    [
        ("Absent This Week",            "Active devotees who did not attend this session"),
        ("Absent 2+ Weeks",             "Devotees missing for two consecutive Sundays"),
        ("Returning Newcomers",         "Devotees who came for the first time after a gap"),
        ("Inactivity Alerts",           "Devotees who haven't attended in 3+ weeks"),
        ("Said Coming — Didn't Come",   "Confirmed \"Yes\" on calling but absent on Sunday"),
    ],
    col_widths=[2.5, 4.5]
)

h3("How to Use the Care Tab")
for step in [
    "Go to the Care tab.",
    "The current session's data loads automatically.",
    "Click on any alert card to see the full list of devotees.",
    "Click a devotee's name to open their profile.",
    "Click \"Export\" to download the list as an Excel file.",
]:
    numbered(step)
tip("Use the \"Said Coming — Didn't Come\" list to talk to coordinators about improving calling accuracy.")

# ── 5H EVENTS ─────────────────────────────────────────────────────────────────
h2("5H. Events")
body("The Events tab lets you create and manage special events (other than regular Sundays).")

h3("Creating an Event")
for step in [
    "Go to the Events tab.",
    "Click \"+ New Event\".",
    "Enter the Event Name, Date, and an optional Description.",
    "Save.",
]:
    numbered(step)

h3("Adding Devotees to an Event")
for step in [
    "Open the event (click its card).",
    "Use the search box to find devotees.",
    "Click \"Add\" next to each devotee's name to register them.",
    "To remove someone, click the \"Remove\" button next to their name.",
]:
    numbered(step)

h3("Exporting Event Attendance")
for step in [
    "Open the event.",
    "Click \"Export\" to download the list of registered devotees as an Excel file.",
]:
    numbered(step)

# ── 5I CALLING MGMT ───────────────────────────────────────────────────────────
h2("5I. Calling Management (Super Admin Only)")
body("The Calling Mgmt tab is only visible to Super Admins. It gives a bird's-eye view of all teams' calling progress across the current week:")
bullet("See which teams have submitted calling")
bullet("See which coordinators are still pending")
bullet("Track submission times and identify late submissions")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — FILTER BAR
# ══════════════════════════════════════════════════════════════════════════════
h1("6. The Filter Bar — Controlling What You See")
body("The Filter Bar (the row of three buttons just below the tabs) is the most important control in the app. Changing these filters changes what every tab shows.")
screenshot("Filter Bar with Session, Team, and Calling By chips")

h3("Session Filter")
bullet("Shows the date of the current Sunday session.")
bullet("Click it to open a dropdown of all past sessions.")
bullet("Changing it shows data for that past session — useful for reviewing old attendance or calling data.")
bullet("Click the ✕ on the chip to go back to the latest session.")

h3("Team Filter")
bullet("Shows \"All Teams\" by default (for Super Admins).")
bullet("Coordinators are locked to their own team and cannot change this.")
bullet("Super Admins can click to filter to one team.")

h3("Calling By Filter")
bullet("Shows \"All Callers\" by default.")
bullet("Click to narrow the view to a specific coordinator's list.")

h3("The Caption Line")
body("Below the three chips, a line of text tells you exactly what you're looking at, for example:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
r = p.add_run("\"Showing all teams, called by Radhika, for Sun 27 Apr 2026\"")
r.italic = True
r.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — ADMIN PANEL
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Admin Panel — Managing Users & Sessions")
body("The Admin Panel is only accessible to Super Admins. Look for the settings or admin icon in the top-right corner of the screen.")

h2("Approving New User Signups")
body("When someone signs up, a badge with a number appears on the admin icon.")
for step in [
    "Click the admin icon → \"Sign-up Requests\".",
    "You'll see a list of pending requests with the name, email, and when they requested access.",
    "For each request, select their Role (Facilitator / Coordinator / Super Admin) and Team.",
    "Click \"Approve\" (green) or \"Reject\" (red).",
]:
    numbered(step)
warn("Until you approve a signup, the user cannot log in — they will remain on the \"Awaiting Approval\" screen.")

h2("Managing Existing Users")
for step in [
    "Admin icon → \"User Management\".",
    "Search for a user by name.",
    "Click their row to open their settings — change Role, Team, Position, or enable Att. Seva.",
    "To remove a user, click the \"Remove\" button (use with caution).",
]:
    numbered(step)

h2("Configuring a Sunday Session")
for step in [
    "Admin icon → \"Session Configuration\".",
    "Fill in the Calling Date, Attendance Date, Topic and Speaker (optional), and Session Type (Regular or Festival).",
    "Click \"Save Session\".",
]:
    numbered(step)

h2("Editing Past Sessions")
for step in [
    "Admin icon → \"Session Management\".",
    "A table shows all past sessions with dates and attendance counts.",
    "Click \"Edit\" on any row to update the topic or mark the session as Cancelled.",
]:
    numbered(step)
body("Note: Marking a session as cancelled does not delete attendance records.")

h2("Clearing Data (Danger Zone)")
for step in [
    "Admin icon → \"Clear Data\".",
    "Three options: Clear by Date, Clear by Team + Date, or Clear ALL Data (requires typing \"DELETE ALL\" to confirm).",
]:
    numbered(step)
warn("Clearing data is permanent and cannot be undone. Always double-check before confirming.")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — EXCEL
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Excel Import & Export")

h2("Exporting")
body("Most lists and reports have an \"Export\" or \"Export Excel\" button. Click it to download the data as a formatted Excel file. Currency is shown in Indian Rupees (₹) and dates are in DD-MM-YYYY format.")

h2("Importing Devotees")
for step in [
    "Go to the Devotees tab.",
    "Click the \"Import\" button.",
    "Select your Excel file.",
    "The app will match columns, detect possible duplicates (by name + mobile), and show a confirmation screen before writing anything.",
    "Review the list and confirm.",
]:
    numbered(step)
warn("Duplicates are detected by name + mobile number together. If a devotee exists with the same name but a different mobile number, they will be treated as a different person.")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — TIPS, WARNINGS & FAQs
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Tips, Warnings & Common Questions")

h2("Quick-Reference Tips")
tips_list = [
    ("Use the Filter Bar first",
     "Before doing any data entry, check that Session, Team, and Calling By filters are set correctly. Everything you enter will be filed under these filters."),
    ("Mobile-friendly",
     "The app works on your phone browser. Open it on Chrome or Safari and tap \"Add to Home Screen\" to use it like an app."),
    ("Profile completeness",
     "The coloured circle on a devotee's card shows how complete their profile is. Green = most fields filled. Try to fill all fields when first adding a devotee."),
    ("Calling before 9 PM",
     "Always submit your calling by 9 PM on the deadline day. Late submissions are flagged in reports."),
    ("Undo attendance quickly",
     "If you marked the wrong person as present, tap \"Undo\" immediately."),
]
for title, text in tips_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"💡  {title}: ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

h2("Technical Warnings")
warnings_list = [
    ("Don't use the browser's back button",
     "This app is a single page. Use the tabs at the top to navigate, not the browser back button."),
    ("Calling submission is time-stamped",
     "Once you click \"Submit Calling,\" the submission time is locked. Reports use this time to check for late submissions."),
    ("Team names must match exactly on import",
     "When importing data from Excel, team names must match exactly (e.g., \"Vishakha\" not \"Visakha\")."),
    ("Cleared data cannot be recovered",
     "The \"Clear Data\" option permanently deletes records. There is no undo."),
    ("Profile images have a 50 KB size limit",
     "When uploading a profile photo in \"Edit Profile,\" keep the file under 50 KB."),
]
for title, text in warnings_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"⚠️  {title}: ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

h2("Common Questions")
faqs = [
    ("I can't see some tabs — is something wrong?",
     "No. The tabs you see depend on your role. Coordinators and Facilitators have fewer tabs than Super Admins. This is by design."),
    ("I updated a calling status but it still shows the old value — is it saved?",
     "Changes are saved as soon as you make them — no \"Save\" button is needed on individual calling rows. If you see the old value, try switching tabs and coming back."),
    ("A devotee isn't showing up in my calling list — why?",
     "Check that the Session and Team filters are correct. Also check if the devotee is marked \"Inactive\" or \"Not Interested\" — those are excluded from the calling list."),
    ("The app is showing old data even after I made changes — what should I do?",
     "Try a hard refresh: Ctrl + Shift + R on Windows, or Cmd + Shift + R on Mac. This forces the app to reload with the latest data."),
    ("Can two people use the app at the same time?",
     "Yes. Multiple people can be logged in simultaneously. If two people update the same devotee at the same time, the last save wins."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"Q: {q}")
    r1.bold = True; r1.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(f"A: {a}")
    r2.font.size = Pt(11)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("This guide was generated from the Sakhi Sang Attendance Tracker codebase. For technical issues, contact your Super Admin.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"d:\riya personal\RAPP\Sadhana Tracker\Sadhana Tracker- Coordinators Position based based\Sakhi-Sang-Attendence\USER_GUIDE.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
